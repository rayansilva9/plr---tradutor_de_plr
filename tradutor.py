import os
import tkinter
import customtkinter
from googletrans import Translator
from docx import Document
from tkinter import filedialog
import aspose.words as aw
from tkinter.ttk import Progressbar
import threading


janela = customtkinter.CTk()
janela.geometry('800x600')
janela.title('Tradutor de PLR')

tradutor = Translator(service_urls=['translate.google.com'])

arquivos = []
arquivos_to_remove = arquivos
plr_txt = []
plr_doc = []

local_de_salvamento = ''

traduzir_para = ['pt', 'fr', 'en']
progress = Progressbar(janela, orient="horizontal",
                       length=100, mode='indeterminate')


def bar():

    if len(arquivos_to_remove) > 0:
        progress.pack(pady=10)
        while len(arquivos_to_remove) > 0:
            import time
            progress['value'] = 20
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 40
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 50
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 60
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 80
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 100
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 80
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 60
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 50
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 40
            janela.update_idletasks()
            time.sleep(0.5)

            progress['value'] = 20
            janela.update_idletasks()
            time.sleep(0.5)
            progress['value'] = 0
    else:
        progress.destroy()


def selectDir():
    global local_de_salvamento
    local_de_salvamento = filedialog.askdirectory()


plrList = tkinter.Listbox(
    janela, width=75, height=10, highlightcolor='white', background='#2b2b2b')

plrList.place(x=170, y=300,)


def obterArquivos():
    global arquivos_to_remove
    arquivos.clear()
    plr_txt.clear()
    plr_doc.clear()
    for i in filedialog.askopenfilenames():
        arquivos.append(i)

    for arquivo in arquivos:
        nome, extensao = os.path.splitext(arquivo)
        plrList.insert(0, os.path.basename(nome))

        if extensao == '.txt':
            plr_txt.append(
                {'nome': os.path.basename(nome), 'path': arquivo, })

        if extensao == '.doc' or extensao == '.docx':
            plr_doc.append(
                {'nome': os.path.basename(nome), 'path': arquivo, })
    arquivos_to_remove = arquivos


def dividir_texto(string, tamanho):
    return [string[i:i+tamanho] for i in range(0, len(string), tamanho)]


def traduzir_txt3():
    for plr in plr_txt:
        for lang in traduzir_para:
            if os.path.exists(f"{local_de_salvamento}/{lang}"):
                pass
            else:
                os.makedirs(f"{local_de_salvamento}/{lang}")
            texto_completo1 = []
            texto_completo2 = ''

            nome_trad = tradutor.translate(plr["nome"].replace(
                "_", " "), src='auto', dest=lang,).text

            with open(plr['path'], "r", encoding='windows-1252') as arquivo:
                partes = dividir_texto(arquivo.read(), 4999)
                for parte in partes:
                    texto_completo1.append(tradutor.translate(
                        parte, src='auto', dest=lang,).text)
                texto_completo2 = '\n'.join(texto_completo1)
                arquivo.close()

            with open(local_de_salvamento + '/' + lang + "/" + nome_trad + '.txt', 'w', encoding='utf-8') as arquivo:
                arquivo.write(texto_completo2)
                arquivo.close()
        arquivos_to_remove.remove(plr["path"])
        bar()


def traduzir_doc3():

    for plr in plr_doc:
        for lang in traduzir_para:
            if os.path.exists(f"{local_de_salvamento}/{lang}"):
                pass
            else:
                os.makedirs(f"{local_de_salvamento}/{lang}")
            texto1 = []
            texto2 = ''
            texto_completo1 = []
            texto_completo2 = ''
            documentTraduzido = Document()
            doc = aw.Document(plr['path'])
            nome_trad = tradutor.translate(
                plr["nome"], src='auto', dest=lang,).text
            doc.save(f"{local_de_salvamento}/{lang}/{nome_trad + '.docx'}")
            documentDocx = Document(
                f"{local_de_salvamento}/{lang}/{nome_trad + '.docx'}")
            for paragraph in documentDocx.paragraphs[1:]:
                texto1.append(paragraph.text)

            texto2 = '\n'.join(texto1)
            partes = dividir_texto(texto2, 4999)

            for parte in partes:
                texto_completo1.append(tradutor.translate(
                    parte, src='auto', dest=lang,).text)
            texto_completo2 = '\n'.join(texto_completo1)
            documentTraduzido.add_paragraph(texto_completo2)
            documentTraduzido.save(
                local_de_salvamento + '/' + lang + "/" + nome_trad + '.docx')
        arquivos_to_remove.remove(plr["path"])
        bar()


def action():
    if len(arquivos) == 0:
        tkinter.messagebox.showinfo("0 arquivo", "Nenhum arquivo selecionado")
    if local_de_salvamento == '':
        tkinter.messagebox.showinfo("", "Selecione onde salvar")
    else:
        threading.Thread(target=traduzir_doc3).start()
        threading.Thread(target=traduzir_txt3).start()
        threading.Thread(target=bar).start()


texto1 = customtkinter.CTkLabel(janela, text='Rodar codigo',)
texto1.pack()

botaoRun = customtkinter.CTkButton(
    janela, text='Selecionar arquivos', command=obterArquivos,)

botaoRun.pack(padx=20, pady=20,)

plrList = tkinter.Listbox(
    janela, width=75, height=10, highlightcolor='white', background='#2b2b2b', border=0, borderwidth=0, )
plrList.place(x=170, y=300,)

botaoSelectDir = customtkinter.CTkButton(janela, text='Salvar em...',
                                         command=selectDir,)
botaoSelectDir.pack(padx=20, pady=20,)

botaoTraduzir = customtkinter.CTkButton(janela, text='Traduzir',
                                        command=action,)
botaoTraduzir.pack(padx=20, pady=20,)

janela.mainloop()
